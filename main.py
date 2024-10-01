import streamlit as st

import pandas as pd
from fuzzywuzzy import fuzz
import phonetics
import jellyfish
import openpyxl
import Levenshtein
import openai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Pt, Inches
import io
import os 

openai_api_key = 


# Funciones de similitud
def soundex_similarity(name1, name2):
    return fuzz.ratio(jellyfish.soundex(name1.lower()), jellyfish.soundex(name2.lower()))

def phonex_similarity(name1, name2):
    return fuzz.ratio(jellyfish.nysiis(name1.lower()), jellyfish.nysiis(name2.lower()))

def levenshtein_similarity(name1, name2):
    return fuzz.ratio(name1.lower(), name2.lower())

def ngram_similarity(s1, s2, n=2):
    s1, s2 = s1.lower().replace(" ", ""), s2.lower().replace(" ", "")
    def get_ngrams(string, n):
        return {string[i:i+n]: 1 for i in range(len(string) - n + 1)}
    ngrams1 = get_ngrams(s1, n)
    ngrams2 = get_ngrams(s2, n)
    common_ngrams = set(ngrams1.keys()) & set(ngrams2.keys())
    similarity = 2 * len(common_ngrams) / (len(ngrams1) + len(ngrams2)) if ngrams1 or ngrams2 else 0.0
    len_diff = abs(len(s1) - len(s2)) / max(len(s1), len(s2)) if max(len(s1), len(s2)) > 0 else 0
    adjusted_similarity = similarity * (1 - len_diff)
    return adjusted_similarity * 100

def phonetic_combined_similarity(name1, name2):
    phonetic_similarity = fuzz.ratio(phonetics.metaphone(name1.lower()), phonetics.metaphone(name2.lower()))
    soundex_sim = soundex_similarity(name1, name2)
    phonex_sim = phonex_similarity(name1, name2)
    return (phonetic_similarity + soundex_sim + phonex_sim) / 3

def orthographic_combined_similarity(name1, name2):
    ortho_similarity = levenshtein_similarity(name1, name2)
    ngram_sim = ngram_similarity(name1, name2)
    return (ortho_similarity + ngram_sim) / 2

def average_similarity(name1, name2):
    phonetic_avg = phonetic_combined_similarity(name1, name2)
    orthographic_avg = orthographic_combined_similarity(name1, name2)
    return (phonetic_avg + orthographic_avg) / 2

def detailed_similarity(name1, name2):
    ortho_similarity = levenshtein_similarity(name1, name2)
    phonetic_similarity = fuzz.ratio(phonetics.metaphone(name1.lower()), phonetics.metaphone(name2.lower()))
    soundex_sim = soundex_similarity(name1, name2)
    phonex_sim = phonex_similarity(name1, name2)
    ngram_sim = ngram_similarity(name1, name2)
    combined_phonetic = phonetic_combined_similarity(name1, name2)
    combined_orthographic = orthographic_combined_similarity(name1, name2)
    avg_similarity = (combined_phonetic + combined_orthographic) / 2
    justificacion = justificar_similitud(name1, name2)
    return {
        "ortho_similarity": ortho_similarity,
        "phonetic_similarity": phonetic_similarity,
        "soundex_sim": soundex_sim,
        "phonex_sim": phonex_sim,
        "ngram_sim": ngram_sim,
        "combined_phonetic": combined_phonetic,
        "combined_orthographic": combined_orthographic,
        "avg_similarity": avg_similarity,
        "justificacion": justificacion
    }

def justificar_similitud(name1, name2):
    ops = Levenshtein.editops(name1.lower(), name2.lower())
    justificacion = []
    for op in ops:
        if op[0] == 'replace':
            justificacion.append(f"Sustituir '{name1[op[1]]}' por '{name2[op[2]]}'")
        elif op[0] == 'insert':
            justificacion.append(f"Insertar '{name2[op[2]]}' en la posición {op[1]}")
        elif op[0] == 'delete':
            justificacion.append(f"Eliminar '{name1[op[1]]}' de la posición {op[1]}")
    return justificacion

def exportar_a_word(nombre_generado, top_5_similitudes, razones):
    doc = Document()
    doc.add_heading("Informe de Similitud", 0)

    # Establecer el tamaño de la fuente y el estilo
    font_size = Pt(11)  # Tamaño de fuente 11
    font_name = 'Times New Roman'  # Nombre de la fuente

    # Añadir el párrafo con formato
    paragraph = doc.add_paragraph(razones.replace('*', '').replace('#', '').replace('-', ''))  # Eliminar asteriscos, signos de número y guiones
    run = paragraph.runs[0]  # Obtener el primer run del párrafo
    run.font.size = font_size  # Establecer el tamaño de la fuente
    run.font.name = font_name  # Establecer el nombre de la fuente
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinear a la izquierda

    # Añadir una sección para la tabla de similitudes en disposición horizontal
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE

    # Establecer márgenes
    section.left_margin = Inches(1)  # Margen izquierdo
    section.right_margin = Inches(1)  # Margen derecho
    section.top_margin = Inches(1)  # Margen superior
    section.bottom_margin = Inches(1)  # Margen inferior

    # Añadir la tabla de similitudes
    doc.add_heading('Tabla de Similitudes', level=1)
    table = doc.add_table(rows=1, cols=9)
    table.autofit = True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Comparación'
    similarity_types = ["ortho_similarity", "phonetic_similarity", "soundex_sim", "phonex_sim", "ngram_sim", "combined_phonetic", "combined_orthographic", "avg_similarity"]
    for i, sim_type in enumerate(similarity_types, start=1):
        hdr_cells[i].text = sim_type

    for nombre_ema, _ in top_5_similitudes:
        row_cells = table.add_row().cells
        comparacion = f"{nombre_generado} vs {nombre_ema}"
        row_cells[0].text = comparacion
        similitud = detailed_similarity(nombre_generado, nombre_ema)
        for i, sim_type in enumerate(similarity_types, start=1):
            row_cells[i].text = f"{similitud[sim_type]:.2f}%"

    # Guardar el documento en un objeto BytesIO
    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)  # Volver al inicio del objeto BytesIO
    return byte_io

def procesar_nombre(nombre_generado, nombres_ema):
    similitudes_lista = []
    
    # Calcular similitudes con todos los nombres de la EMA
    for nombre_ema in nombres_ema:
        avg_sim = average_similarity(nombre_generado, nombre_ema)
        similitudes_lista.append((nombre_ema, avg_sim))

    # Ordenar y seleccionar los 5 nombres EMA más similares
    similitudes_lista.sort(key=lambda x: x[1], reverse=True)
    top_5_similitudes = similitudes_lista[:5]

    st.write("Los 5 nombres con mayor similitud media son:")
    for nombre_ema, avg_sim in top_5_similitudes:
        st.write(f"{nombre_ema} con una similitud media de {avg_sim:.2f}%")

    # Elegir el nombre EMA con mayor similitud al nombre generado
    nombre_riesgo, similitud_riesgo = top_5_similitudes[0]

    ver_detalle = st.radio("¿Quieres ver los detalles?", ("Sí", "No"), index=1)
    if ver_detalle == 'Sí':
        for nombre_ema, _ in top_5_similitudes:
            similitud = detailed_similarity(nombre_generado, nombre_ema)
            st.write(f"\nAnálisis detallado para '{nombre_generado} vs {nombre_ema}':")
            for key, value in similitud.items():
                if key != 'justificacion':
                    st.write(f"{key}: {value:.2f}%")

        # Llamada a la API de ChatGPT para evaluar el riesgo
        prompt = (
            f"El nombre '{nombre_generado}' presenta una similitud de {similitud_riesgo:.2f}% con {nombre_riesgo}. "
            "¿Cuáles son las razones por las que la EMA podría conceder o no este nombre teniendo en cuenta que con una similitud del 50% o más, es probable que no lo conceda? "
            "Considera la importancia de la similitud fonética y ortográfica. "
            "Actúa como un evaluador oficial de la European Medicines Agency (EMA). "
            "Quiero que evalúes el nombre generado '{nombre_generado}' en comparación con '{nombre_riesgo}' y determines si se parece demasiado a este nombre, "
            "hasta el punto de que un usuario medio, tanto en español como en inglés, podría confundirse. "
            "Utiliza las funciones del código para realizar esta evaluación, comparando no solo la similitud literal de los nombres, "
            "sino también aspectos fonéticos y de pronunciación en ambos idiomas. "
            "El objetivo es determinar si este nuevo nombre es registrable, o si presenta un riesgo de confusión con el nombre seleccionado. "
            "Por favor, ofrece una explicación articulada, razonando los motivos por los cuales el nuevo nombre podría generar confusión desde un punto de vista global, "
            "basándote en las reglas de la EMA. También explora si hay diferencias en cómo los usuarios de habla inglesa o española podrían percibir la similitud. "
            "Finalmente, toma una decisión razonada sobre si el nombre propuesto debería ser aceptado o rechazado en base a las similitudes identificadas."
            "Pon más énfasis en la similitud que ocurre al principio del nombre, en particular en las tres primeras letras del nombre y en identidad en de letras."
        )

        try:
            response = openai.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Eres un examinador oficial de la European Medicines Agency (EMA). Tu misión es evaluar si un nombre puede registrarse o si, por el contrario, no puede al parecerse demasiado (más de 50%) a otros nombres ya registrados."},
                    {"role": "user", "content": prompt}
                ]
            )

            if response:
                razones = response.choices[0].message['content'].strip()
                st.write(razones)

                # Opción para exportar a Word
                byte_io = exportar_a_word(nombre_generado, top_5_similitudes, razones)
                st.download_button("Descargar Informe en Word", byte_io, "informe_similitud.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            else:
                st.write("Error al obtener la respuesta de la API.")
        except Exception as e:
            st.write(f"Error al obtener la respuesta de la API: {str(e)}")

    return top_5_similitudes

def main():
    st.title("EMA_Bot")

    if "messages" not in st.session_state:
        st.session_state.messages = []

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    try:
        st.write("Leyendo el archivo EMA.xlsx...")
        nombres_ema = pd.read_excel('EMA.xlsx')['Nombre'].tolist()
        st.write("Archivo leído correctamente.")
    except Exception as e:
        st.write(f"Error al leer el archivo Excel: {e}")
        return

    nombre_generado = st.text_input("Introduce un nombre para analizar:")
    if nombre_generado:
        st.write(f"Analizando el nombre: {nombre_generado}")
        procesar_nombre(nombre_generado, nombres_ema)

if __name__ == "__main__":
    main()
